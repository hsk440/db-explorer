import streamlit as st
import psycopg2
import psycopg2.extras
import pandas as pd
import plotly.express as px
import os
import json
import re
import logging
import time
import io
from datetime import datetime
from pathlib import Path

import skills
import llm_client

# ---------------------------------------------------------------------------
# Logging setup
# ---------------------------------------------------------------------------
LOG_DIR = Path(__file__).parent / "logs"
LOG_DIR.mkdir(exist_ok=True)
LOG_FILE = LOG_DIR / f"app_{datetime.now().strftime('%Y-%m-%d')}.log"

logger = logging.getLogger("db_explorer")
if not logger.handlers:
    logger.setLevel(logging.DEBUG)
    fmt = logging.Formatter(
        "%(asctime)s | %(levelname)-7s | %(message)s", datefmt="%Y-%m-%d %H:%M:%S"
    )
    # File handler - detailed logs
    fh = logging.FileHandler(LOG_FILE, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(fmt)
    logger.addHandler(fh)
    # Console handler - errors only
    ch = logging.StreamHandler()
    ch.setLevel(logging.ERROR)
    ch.setFormatter(fmt)
    logger.addHandler(ch)

logger.info("=" * 60)
logger.info("App started")

# ---------------------------------------------------------------------------
# Page config
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="DB Explorer",
    page_icon="🗄️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

SAFE_STARTS = re.compile(
    r"^\s*(SELECT|WITH|EXPLAIN|SHOW|VALUES)\b",
    re.IGNORECASE,
)


def is_read_only_query(sql: str) -> bool:
    """Check that the query starts with a safe read-only keyword.
    The DB session is already read-only as the real safety net,
    so this just catches obviously wrong queries."""
    cleaned = re.sub(r"--.*$", "", sql, flags=re.MULTILINE)
    cleaned = re.sub(r"/\*.*?\*/", "", cleaned, flags=re.DOTALL)
    cleaned = cleaned.strip()
    is_safe = bool(SAFE_STARTS.match(cleaned))
    if not is_safe:
        logger.warning(f"BLOCKED QUERY | starts with: {cleaned[:80]}")
    return is_safe


def get_connection(host, port, dbname, user, password):
    """Create a READ-ONLY connection to PostgreSQL."""
    t0 = time.time()
    try:
        conn = psycopg2.connect(
            host=host,
            port=port,
            dbname=dbname,
            user=user,
            password=password,
            options="-c default_transaction_read_only=on",
            connect_timeout=10,
        )
        conn.set_session(readonly=True, autocommit=True)
        logger.info(f"DB connected to {user}@{host}:{port}/{dbname} ({time.time()-t0:.2f}s)")
        return conn
    except Exception as e:
        logger.error(f"DB connection FAILED to {user}@{host}:{port}/{dbname} - {e}")
        raise


@st.cache_data(ttl=300, show_spinner=False)
def fetch_schemas(_conn_id, host, port, dbname, user, password):
    conn = get_connection(host, port, dbname, user, password)
    cur = conn.cursor()
    cur.execute(
        "SELECT schema_name FROM information_schema.schemata "
        "WHERE schema_name NOT IN ('pg_catalog','information_schema','pg_toast') "
        "ORDER BY schema_name"
    )
    schemas = [r[0] for r in cur.fetchall()]
    cur.close()
    conn.close()
    return schemas


@st.cache_data(ttl=300, show_spinner=False)
def fetch_tables(_conn_id, host, port, dbname, user, password, schema):
    conn = get_connection(host, port, dbname, user, password)
    cur = conn.cursor()
    cur.execute(
        "SELECT table_name, table_type FROM information_schema.tables "
        "WHERE table_schema = %s ORDER BY table_name",
        (schema,),
    )
    tables = cur.fetchall()
    cur.close()
    conn.close()
    return tables


@st.cache_data(ttl=300, show_spinner=False)
def fetch_columns(_conn_id, host, port, dbname, user, password, schema, table):
    conn = get_connection(host, port, dbname, user, password)
    cur = conn.cursor()
    cur.execute(
        "SELECT column_name, data_type, is_nullable, column_default "
        "FROM information_schema.columns "
        "WHERE table_schema = %s AND table_name = %s "
        "ORDER BY ordinal_position",
        (schema, table),
    )
    cols = cur.fetchall()
    cur.close()
    conn.close()
    return cols


@st.cache_data(ttl=300, show_spinner=False)
def fetch_row_count(_conn_id, host, port, dbname, user, password, schema, table):
    conn = get_connection(host, port, dbname, user, password)
    cur = conn.cursor()
    cur.execute(f'SELECT count(*) FROM "{schema}"."{table}"')
    count = cur.fetchone()[0]
    cur.close()
    conn.close()
    return count


def run_query(host, port, dbname, user, password, sql, limit=500):
    logger.info(f"QUERY START | limit={limit} | {sql[:200]}")
    t0 = time.time()
    try:
        conn = get_connection(host, port, dbname, user, password)
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(sql)
        rows = cur.fetchmany(limit)
        columns = [desc[0] for desc in cur.description] if cur.description else []
        cur.close()
        conn.close()
        elapsed = time.time() - t0
        row_count = len(rows) if rows else 0
        logger.info(f"QUERY OK | {row_count} rows | {elapsed:.2f}s | {sql[:120]}")
        return pd.DataFrame(rows, columns=columns) if rows else pd.DataFrame()
    except Exception as e:
        logger.error(f"QUERY FAILED | {time.time()-t0:.2f}s | {e} | {sql[:200]}")
        raise


def get_schema_summary(host, port, dbname, user, password):
    """Build a compact schema description for the AI."""
    conn = get_connection(host, port, dbname, user, password)
    cur = conn.cursor()
    cur.execute(
        """
        SELECT t.table_schema, t.table_name, c.column_name, c.data_type
        FROM information_schema.tables t
        JOIN information_schema.columns c
          ON t.table_schema = c.table_schema AND t.table_name = c.table_name
        WHERE t.table_schema NOT IN ('pg_catalog','information_schema','pg_toast')
          AND t.table_type = 'BASE TABLE'
        ORDER BY t.table_schema, t.table_name, c.ordinal_position
        """
    )
    rows = cur.fetchall()
    cur.close()
    conn.close()

    tables = {}
    for schema, table, col, dtype in rows:
        key = f"{schema}.{table}"
        tables.setdefault(key, []).append(f"  {col} ({dtype})")

    parts = []
    for tbl, cols in tables.items():
        parts.append(f"{tbl}:\n" + "\n".join(cols))
    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# Export helpers: CSV, Excel, Word, Text
# ---------------------------------------------------------------------------

def df_to_csv_bytes(df: pd.DataFrame) -> bytes:
    """Convert DataFrame to CSV bytes."""
    return df.to_csv(index=False).encode("utf-8")


def df_to_excel_bytes(df: pd.DataFrame, sheet_name: str = "Data") -> bytes:
    """Convert DataFrame to Excel (.xlsx) bytes."""
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name=sheet_name[:31], index=False)
    return buf.getvalue()


def build_text_report(question: str, analysis: str, sql: str, df: pd.DataFrame) -> bytes:
    """Build a plain text report with the question, analysis, SQL, and data."""
    lines = []
    lines.append("=" * 70)
    lines.append("DB Explorer - Analysis Report")
    lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("=" * 70)
    lines.append("")
    lines.append(f"QUESTION:")
    lines.append(question)
    lines.append("")
    lines.append("ANALYSIS:")
    lines.append("-" * 70)
    lines.append(analysis)
    lines.append("")
    lines.append("SQL QUERY USED:")
    lines.append("-" * 70)
    lines.append(sql)
    lines.append("")
    lines.append(f"DATA ({len(df)} rows):")
    lines.append("-" * 70)
    if df.empty:
        lines.append("(no rows returned)")
    else:
        lines.append(df.to_string(index=False, max_rows=1000))
    lines.append("")
    return "\n".join(lines).encode("utf-8")


def build_docx_report(question: str, analysis: str, sql: str, df: pd.DataFrame) -> bytes:
    """Build a Word (.docx) report with formatting."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title
    title = doc.add_heading("DB Explorer - Analysis Report", level=0)

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True

    # Question
    doc.add_heading("Question", level=1)
    doc.add_paragraph(question)

    # Analysis - handle basic markdown (bold, bullets)
    doc.add_heading("Analysis", level=1)
    for line in analysis.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_markdown_runs(p, line[2:])
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        else:
            p = doc.add_paragraph()
            _add_markdown_runs(p, line)

    # SQL
    doc.add_heading("SQL Query Used", level=1)
    sql_para = doc.add_paragraph()
    sql_run = sql_para.add_run(sql)
    sql_run.font.name = "Courier New"
    sql_run.font.size = Pt(9)

    # Data table
    doc.add_heading(f"Data ({len(df)} rows)", level=1)
    if df.empty:
        doc.add_paragraph("(no rows returned)")
    else:
        # Cap at 100 rows to keep docx manageable
        display_df = df.head(100)
        table = doc.add_table(rows=1, cols=len(display_df.columns))
        table.style = "Light Grid Accent 1"

        # Header row
        hdr = table.rows[0].cells
        for i, col in enumerate(display_df.columns):
            hdr[i].text = str(col)
            for para in hdr[i].paragraphs:
                for run in para.runs:
                    run.bold = True

        # Data rows
        for _, row in display_df.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = "" if pd.isna(val) else str(val)

        if len(df) > 100:
            doc.add_paragraph(f"(showing first 100 of {len(df)} rows)").italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_markdown_runs(paragraph, text: str):
    """Add text to a paragraph, handling **bold** markdown."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


TRUNCATION_NOTICE = "\n\n_[Response was truncated — ask a narrower question for full detail.]_"


def _build_chat_messages(chat_history, question, use_summary=False):
    """Build OpenAI-style message list from chat history + current question."""
    messages = []
    if chat_history:
        for msg in chat_history:
            if msg["role"] == "user":
                messages.append({"role": "user", "content": msg["content"]})
            elif msg["role"] == "assistant":
                if use_summary:
                    content = msg.get("summary") or msg.get("content", "")
                else:
                    content = msg.get("content", "")
                if content:
                    messages.append({"role": "assistant", "content": content})
    messages.append({"role": "user", "content": question})
    return messages


def ask_claude_for_sql(question, schema_text, api_key, chat_history=None, model=None, tier="standard"):
    """Route the user's question. Returns (sql_or_sentinel, error)."""
    if model is None:
        model = llm_client.smart_pick_model("anthropic", "sql")

    system_prompt = f"""You are a PostgreSQL router/assistant. Given the database schema below, decide how to handle the user's question and respond.

YOUR THREE RESPONSE MODES (return EXACTLY one):

MODE 1 — Simple SQL (default): Return a single SELECT query that answers the question.
- ONLY SELECT/WITH/EXPLAIN/SHOW/VALUES. Never write INSERT, UPDATE, DELETE, DROP, or any data-modifying statement.
- Return ONLY the SQL query, no explanation, no markdown fences.
- Use double quotes for identifiers with special characters.
- Limit results to 500 rows unless the user asks for more.
- Prefer this mode when the answer is one query + one explanation.

MODE 2 — Agent (multi-step with tools): Return EXACTLY the token: USE_AGENT_MODE
- Use this when the user asks for:
  * An Excel/spreadsheet file (.xlsx), Word/doc file (.docx), or structured report
  * Multiple organized queries (e.g. "break down by category", "for each area show...")
  * A file output with organized sections or sheets
  * Anything that needs multiple queries combined into a single output
- The agent has tools to run multiple queries, create Excel/Word files, and return them for download.

MODE 3 — Conversational (no DB needed): Return EXACTLY the token: NO_SQL_NEEDED
- Use this for: thanks, greetings, explanations of concepts, "what does this mean", schema overview questions that can be answered from the schema text above.

DATABASE SCHEMA:
{schema_text}"""

    messages = _build_chat_messages(chat_history, question, use_summary=True)

    logger.info(f"AI SQL REQUEST | model={model} | question: {question[:150]}")
    try:
        resp = llm_client.llm_complete(
            model=model, messages=messages, api_key=api_key,
            system=system_prompt, max_tokens=1024, tier=tier,
        )
    except Exception as e:
        logger.error(f"AI SQL FAILED | {e}")
        return None, str(e)

    sql = resp.text
    if not sql:
        logger.warning("AI SQL returned empty text")
        return None, "AI returned an empty response"

    # Strip markdown fences case-insensitively
    sql = re.sub(r"^```(?:sql|postgresql)?\s*\n?", "", sql, flags=re.IGNORECASE)
    sql = re.sub(r"\n?```\s*$", "", sql)
    sql = sql.strip()

    # Handle stop reasons
    if resp.stop_reason == llm_client.STOP_MAX_TOKENS:
        logger.warning(f"AI SQL hit max_tokens | partial={sql[:120]}")
        # For the routing call, truncation usually means the model started explaining —
        # try to recover a valid sentinel if present
        if "USE_AGENT_MODE" in sql.upper():
            return "USE_AGENT_MODE", None
        if "NO_SQL_NEEDED" in sql.upper():
            return "NO_SQL_NEEDED", None
        # Otherwise return what we got, caller will route it
    elif resp.stop_reason == llm_client.STOP_REFUSAL:
        return None, "The model declined to answer. Try rephrasing your question."

    logger.info(f"AI SQL OK | sql: {sql[:150]}")
    return sql, None


def ask_claude_to_analyze(question, schema_text, data_csv, sql_used, api_key,
                          chat_history=None, model=None, tier="standard"):
    """Analyze query results and produce a conversational answer."""
    if model is None:
        model = llm_client.smart_pick_model("anthropic", "analyze")

    system_prompt = f"""You are a friendly, intelligent data analyst assistant. The user is non-technical.
You help them understand their PostgreSQL database by answering questions in plain English.

YOUR ROLE:
- Analyze the data returned from queries and give clear, insightful answers.
- Highlight key findings, trends, outliers, or patterns you notice.
- Use simple language. Avoid jargon.
- When relevant, suggest follow-up questions the user might want to ask.
- If the data is empty or unexpected, explain what that might mean.
- Format numbers nicely (e.g., 1,234 instead of 1234, percentages where useful).
- Use bullet points, bold text, and short paragraphs for readability.
- If you spot something interesting or unusual in the data, proactively point it out.

APP FEATURES YOU CAN MENTION:
- The user can export every answer's data using the buttons below each response:
  * "CSV (data)" — spreadsheet-compatible
  * "Excel (data)" — .xlsx file
  * "Word (full report)" — .docx with your analysis + data
  * "Text (full report)" — .txt with your analysis + data
- If the user asks to save/export/download, tell them to click those buttons (do NOT say you can't create files).

DATABASE SCHEMA:
{schema_text}"""

    user_content = f"""My question: {question}

SQL query used:
{sql_used}

Query results (CSV):
{data_csv}

Please analyze this data and answer my question in a clear, helpful way."""

    messages = _build_chat_messages(chat_history, user_content)

    logger.info(f"AI ANALYZE REQUEST | model={model} | question: {question[:150]}")
    try:
        resp = llm_client.llm_complete(
            model=model, messages=messages, api_key=api_key,
            system=system_prompt, max_tokens=4096, tier=tier,
        )
    except Exception as e:
        logger.error(f"AI ANALYZE FAILED | {e}")
        return None, str(e)

    text = resp.text
    if resp.stop_reason == llm_client.STOP_MAX_TOKENS:
        logger.warning("AI ANALYZE hit max_tokens — appending truncation notice")
        text = (text or "") + TRUNCATION_NOTICE
    elif resp.stop_reason == llm_client.STOP_REFUSAL:
        return None, "The model declined to answer this question."

    return (text or "").strip(), None


def ask_claude_no_sql(question, schema_text, api_key, chat_history=None, model=None, tier="standard"):
    """Handle conversational questions that don't need a SQL query."""
    if model is None:
        model = llm_client.smart_pick_model("anthropic", "conversational")

    system_prompt = f"""You are a friendly, intelligent data analyst assistant. The user is non-technical.
You are chatting about their PostgreSQL database.
Answer their question based on the conversation history and your knowledge of their database schema.
Be helpful, thorough, and suggest what they could ask next if appropriate.

If the user asks to export, download, or save results to a file (CSV, Excel, Word, text, etc.):
- Tell them the app has EXPORT BUTTONS below each answer that has data
- The buttons are: "CSV (data)", "Excel (data)", "Word (full report)", "Text (full report)"
- Do NOT say you can't create files — direct them to the export buttons

Use bullet points, bold text, and headings to organize your response.

DATABASE SCHEMA:
{schema_text}"""

    messages = _build_chat_messages(chat_history, question)

    logger.info(f"AI CHAT REQUEST | model={model} | question: {question[:150]}")
    try:
        resp = llm_client.llm_complete(
            model=model, messages=messages, api_key=api_key,
            system=system_prompt, max_tokens=4096, tier=tier,
        )
    except Exception as e:
        logger.error(f"AI CHAT FAILED | {e}")
        return None, str(e)

    text = resp.text
    if resp.stop_reason == llm_client.STOP_MAX_TOKENS:
        logger.warning("AI CHAT hit max_tokens — appending truncation notice")
        text = (text or "") + TRUNCATION_NOTICE
    elif resp.stop_reason == llm_client.STOP_REFUSAL:
        return None, "The model declined to answer this question."

    return (text or "").strip(), None


# ---------------------------------------------------------------------------
# Agent loop - tool use / skills
# ---------------------------------------------------------------------------

AGENT_MAX_STEPS = 10
AGENT_MAX_TOKENS = 8192


def run_agent_loop(question, schema_text, api_key, conn_kwargs, chat_history=None,
                   progress_cb=None, model=None, tier="standard"):
    """
    Run an agentic loop using tools (skills) to answer a question.
    Supports multi-step queries and artifact generation (Excel / Word).

    Handles all stop reasons gracefully (end_turn, tool_use, max_tokens, refusal, unknown).

    Returns: (final_text, artifacts, error)
    """
    if model is None:
        model = llm_client.smart_pick_model("anthropic", "agent")

    # System prompt WITH prompt caching on the schema (the bulky, stable part)
    # LiteLLM translates Anthropic-style cache_control for Gemini automatically
    system = [
        {
            "type": "text",
            "text": (
                "You are an autonomous data analyst agent for a PostgreSQL database. "
                "The user is non-technical. You have tools to query the database and "
                "create downloadable Excel/Word files.\n\n"
                "WORKFLOW:\n"
                "1. Call `query_database` to gather the data you need (can call multiple times).\n"
                "2. If the user asked for a file, call `create_excel_artifact` or `create_word_artifact` "
                "   to build the actual file. Organize data into clear sheets or sections.\n"
                "3. After creating any artifact, end your turn with a short message telling "
                "   the user what you built and that they can download it below.\n\n"
                "RULES:\n"
                "- Only read-only SELECT/WITH queries. The database is in read-only mode.\n"
                "- If a query fails with an error, fix the SQL and retry — don't give up.\n"
                "- Use table aliases when joining information_schema tables to avoid ambiguous columns.\n"
                "- Be efficient — combine queries where possible rather than running one per row.\n"
                "- Present final results in plain English. Don't dump raw SQL on the user.\n"
                "- If the user asked for an Excel or Word file, you MUST create the artifact before ending.\n"
                "- If the result would produce >200 rows across sheets/sections, SPLIT into multiple "
                "  artifacts (e.g., one file per major category) rather than one giant file. "
                "  Each create_*_artifact call should stay focused to avoid hitting response size limits.\n"
            ),
        },
        {
            "type": "text",
            "text": f"DATABASE SCHEMA:\n{schema_text}",
            "cache_control": {"type": "ephemeral"},  # 5min cache
        },
    ]

    # Build messages from chat history (lean: last 6 turns, use summary when available)
    messages = []
    if chat_history:
        trimmed = chat_history[-6:]
        for msg in trimmed:
            if msg["role"] == "user":
                messages.append({"role": "user", "content": msg["content"]})
            elif msg["role"] == "assistant":
                content = msg.get("summary") or msg.get("content", "")
                if content:
                    messages.append({"role": "assistant", "content": content})
    messages.append({"role": "user", "content": question})

    logger.info(f"AGENT START | model={model} | question: {question[:150]}")
    artifacts = []
    total_in, total_out = 0, 0
    t0 = time.time()
    continuation_used = False  # safety: only 1 max_tokens auto-continuation per run

    for step in range(1, AGENT_MAX_STEPS + 1):
        if progress_cb:
            progress_cb(f"Step {step}: Thinking...")

        try:
            resp = llm_client.llm_complete(
                model=model, messages=messages, api_key=api_key,
                system=system, max_tokens=AGENT_MAX_TOKENS,
                tools=skills.TOOL_DEFINITIONS, tier=tier,
            )
        except Exception as e:
            logger.error(f"AGENT STEP {step} API FAILED | {e}")
            return None, artifacts, str(e)

        total_in += resp.input_tokens
        total_out += resp.output_tokens
        logger.info(
            f"AGENT STEP {step} | stop={resp.stop_reason} | "
            f"tokens in={resp.input_tokens} cached={resp.cached_input_tokens} out={resp.output_tokens}"
        )

        # ---- Handle stop reasons ----

        if resp.stop_reason == llm_client.STOP_END_TURN:
            elapsed = time.time() - t0
            logger.info(
                f"AGENT DONE | {step} steps | {elapsed:.1f}s | "
                f"tokens in={total_in} out={total_out} | artifacts={len(artifacts)}"
            )
            return resp.text.strip() if resp.text else "", artifacts, None

        if resp.stop_reason == llm_client.STOP_TOOL_USE:
            if not resp.tool_calls:
                # Unexpected: tool_use stop but no tool_calls — treat as end_turn
                logger.warning("AGENT stop_reason=tool_use but no tool_calls returned")
                return resp.text.strip() if resp.text else "", artifacts, None

            # Append assistant message with tool calls
            messages.append(llm_client.build_assistant_tool_call_message(resp.text, resp.tool_calls))

            # Execute each tool call and append tool results
            for tc in resp.tool_calls:
                if progress_cb:
                    progress_cb(f"Step {step}: Running {tc.name}...")
                result, artifact = skills.dispatch_tool(
                    tc.name, tc.input,
                    run_query_fn=run_query,
                    is_safe_fn=is_read_only_query,
                    conn_kwargs=conn_kwargs,
                )
                if artifact:
                    artifacts.append(artifact)
                messages.append(llm_client.build_tool_result_message(
                    tool_call_id=tc.id,
                    content=json.dumps(result, default=str),
                ))
            continue

        if resp.stop_reason == llm_client.STOP_MAX_TOKENS:
            # The user's reported bug: don't crash, recover gracefully
            logger.warning(f"AGENT hit max_tokens at step {step}")

            if resp.tool_calls:
                # Incomplete tool_call — typically a huge create_excel_artifact that got truncated
                logger.warning(
                    f"AGENT max_tokens mid-tool-call ({len(resp.tool_calls)} calls, "
                    f"last={resp.tool_calls[-1].name}) — returning partial"
                )
                msg = (
                    "I started building a large response but ran out of space mid-way. "
                    "Could you narrow the request? For example, ask for fewer categories at a time, "
                    "or split into multiple questions (e.g. 'just the billing tables first')."
                )
                if artifacts:
                    msg = f"I built {len(artifacts)} file(s) before running out of space, but couldn't finish the rest. " + msg
                return msg, artifacts, None

            # Text-only max_tokens — try to continue once
            if not continuation_used and resp.text:
                continuation_used = True
                logger.info("AGENT attempting max_tokens auto-continuation")
                messages.append({"role": "assistant", "content": resp.text})
                messages.append({
                    "role": "user",
                    "content": "Your previous response was cut off. Please continue from exactly where you left off.",
                })
                continue

            # Already used continuation or no text — return what we have with a note
            text = (resp.text or "") + TRUNCATION_NOTICE
            return text, artifacts, None

        if resp.stop_reason == llm_client.STOP_REFUSAL:
            logger.warning(f"AGENT refused at step {step}")
            return (
                "The model declined to complete this request. "
                "Try rephrasing, or break it into smaller parts."
            ), artifacts, None

        # Unknown stop reason — treat as end_turn to be safe (don't crash)
        logger.warning(f"AGENT unknown stop_reason={resp.stop_reason} — treating as end_turn")
        return resp.text.strip() if resp.text else "", artifacts, None

    # Exhausted iterations
    logger.warning(f"AGENT exhausted {AGENT_MAX_STEPS} steps")
    return (
        f"I reached my step limit ({AGENT_MAX_STEPS}) without finishing. "
        "Please try breaking your request into smaller parts."
    ), artifacts, None


# ---------------------------------------------------------------------------
# Session state defaults
# ---------------------------------------------------------------------------
for key, default in {
    "connected": False,
    "host": "",
    "port": "5432",
    "dbname": "",
    "user": "",
    "password": "",
    "chat_history": [],
    "artifacts": {},  # artifact_id -> artifact dict (with bytes)
    # LLM provider settings (Claude is default)
    "provider": "anthropic",
    "anthropic_api_key": os.environ.get("ANTHROPIC_API_KEY", ""),
    "gemini_api_key": os.environ.get("GEMINI_API_KEY", ""),
    "model_label": "Auto (smart routing)",  # human label
    "tier": "standard",
    "smart_routing": True,
}.items():
    if key not in st.session_state:
        st.session_state[key] = default


def current_api_key() -> str:
    """Return the API key for the currently selected provider."""
    if st.session_state.provider == "gemini":
        return st.session_state.gemini_api_key
    return st.session_state.anthropic_api_key


def pick_model(task: str) -> str:
    """Return the model id to use for a given task, honoring user override."""
    if st.session_state.smart_routing and st.session_state.model_label == "Auto (smart routing)":
        return llm_client.smart_pick_model(st.session_state.provider, task)
    # User chose a specific model
    catalog = llm_client.ANTHROPIC_MODELS if st.session_state.provider == "anthropic" else llm_client.GEMINI_MODELS
    return catalog.get(st.session_state.model_label, llm_client.smart_pick_model(st.session_state.provider, task))


# ---------------------------------------------------------------------------
# Sidebar - Connection
# ---------------------------------------------------------------------------
with st.sidebar:
    st.header("Database Connection")

    if not st.session_state.connected:
        with st.form("connect_form"):
            host = st.text_input("Host / IP", value=st.session_state.host)
            port = st.text_input("Port", value=st.session_state.port)
            dbname = st.text_input("Database Name", value=st.session_state.dbname)
            user = st.text_input("Username", value=st.session_state.user)
            password = st.text_input("Password", type="password", value=st.session_state.password)
            submitted = st.form_submit_button("Connect (Read-Only)", type="primary")

        if submitted:
            with st.spinner("Connecting..."):
                try:
                    conn = get_connection(host, int(port), dbname, user, password)
                    conn.close()
                    st.session_state.connected = True
                    st.session_state.host = host
                    st.session_state.port = port
                    st.session_state.dbname = dbname
                    st.session_state.user = user
                    st.session_state.password = password
                    st.rerun()
                except Exception as e:
                    st.error(f"Connection failed: {e}")
    else:
        st.success(f"Connected to **{st.session_state.dbname}**")
        st.caption(f"{st.session_state.user}@{st.session_state.host}:{st.session_state.port}")
        st.caption("Mode: **READ-ONLY**")
        if st.button("Disconnect"):
            st.session_state.connected = False
            st.rerun()

    st.divider()
    st.header("AI Settings")

    # Provider
    provider_label = st.selectbox(
        "Provider",
        ["Anthropic (Claude)", "Google (Gemini)"],
        index=0 if st.session_state.provider == "anthropic" else 1,
        help="Anthropic is the default. Switch to Gemini to use Google models.",
    )
    st.session_state.provider = "anthropic" if provider_label.startswith("Anthropic") else "gemini"

    # Model dropdown — options depend on provider
    if st.session_state.provider == "anthropic":
        model_options = ["Auto (smart routing)"] + list(llm_client.ANTHROPIC_MODELS.keys())
    else:
        model_options = ["Auto (smart routing)"] + list(llm_client.GEMINI_MODELS.keys())

    # Reset model_label if current selection isn't in options (e.g. after switching provider)
    if st.session_state.model_label not in model_options:
        st.session_state.model_label = "Auto (smart routing)"

    st.session_state.model_label = st.selectbox(
        "Model",
        model_options,
        index=model_options.index(st.session_state.model_label),
        help="'Auto' picks the best model per task (cheaper for routing, more capable for agent). Select a specific model to override.",
    )
    st.session_state.smart_routing = st.session_state.model_label == "Auto (smart routing)"

    # Tier dropdown — options depend on provider
    tier_options = llm_client.TIERS_BY_PROVIDER[st.session_state.provider]
    if st.session_state.tier not in tier_options:
        st.session_state.tier = "standard"
    st.session_state.tier = st.selectbox(
        "Tier",
        tier_options,
        index=tier_options.index(st.session_state.tier),
        format_func=lambda t: t.capitalize(),
        help=(
            "Standard: default rate. "
            "Priority: faster/guaranteed SLA (higher cost). "
            "Flex: 50% off, opportunistic (Gemini only). "
            "Batch: 50% off, async (Gemini only)."
        ),
    )

    # API key — label depends on provider
    if st.session_state.provider == "anthropic":
        st.session_state.anthropic_api_key = st.text_input(
            "Anthropic API Key",
            value=st.session_state.anthropic_api_key,
            type="password",
            help="Get yours at console.anthropic.com",
        )
    else:
        st.session_state.gemini_api_key = st.text_input(
            "Gemini API Key",
            value=st.session_state.gemini_api_key,
            type="password",
            help="Get yours at aistudio.google.com/apikey",
        )


# ---------------------------------------------------------------------------
# Main area
# ---------------------------------------------------------------------------
def conn_params():
    return dict(
        host=st.session_state.host,
        port=int(st.session_state.port),
        dbname=st.session_state.dbname,
        user=st.session_state.user,
        password=st.session_state.password,
    )


def conn_id():
    s = st.session_state
    return f"{s.host}:{s.port}/{s.dbname}/{s.user}"


if not st.session_state.connected:
    st.title("Database Explorer")
    st.info("Enter your PostgreSQL credentials in the sidebar to get started. The connection is **read-only** so you cannot accidentally modify any data.")
    st.stop()

# ---------------------------------------------------------------------------
# Tabs
# ---------------------------------------------------------------------------
tab_schema, tab_sql, tab_nl, tab_analytics, tab_logs = st.tabs([
    "Schema Explorer",
    "SQL Query",
    "AI Assistant",
    "Quick Analytics",
    "Activity Log",
])

# --- Schema Explorer ---
with tab_schema:
    st.header("Schema Explorer")
    schemas = fetch_schemas(conn_id(), **conn_params())

    if not schemas:
        st.warning("No schemas found.")
    else:
        sel_schema = st.selectbox("Schema", schemas, index=schemas.index("public") if "public" in schemas else 0)
        tables = fetch_tables(conn_id(), **conn_params(), schema=sel_schema)

        if not tables:
            st.info("No tables in this schema.")
        else:
            st.markdown(f"**{len(tables)} tables** in `{sel_schema}`")

            for tbl_name, tbl_type in tables:
                with st.expander(f"{'VIEW' if tbl_type == 'VIEW' else 'TABLE'} &mdash; {tbl_name}"):
                    cols = fetch_columns(conn_id(), **conn_params(), schema=sel_schema, table=tbl_name)
                    col_df = pd.DataFrame(cols, columns=["Column", "Type", "Nullable", "Default"])
                    st.dataframe(col_df, use_container_width=True, hide_index=True)

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button(f"Row count", key=f"cnt_{sel_schema}_{tbl_name}"):
                            count = fetch_row_count(conn_id(), **conn_params(), schema=sel_schema, table=tbl_name)
                            st.metric("Total Rows", f"{count:,}")
                    with col2:
                        if st.button(f"Preview 10 rows", key=f"preview_{sel_schema}_{tbl_name}"):
                            df = run_query(**conn_params(), sql=f'SELECT * FROM "{sel_schema}"."{tbl_name}" LIMIT 10')
                            st.dataframe(df, use_container_width=True, hide_index=True)


# --- SQL Query ---
with tab_sql:
    st.header("SQL Query")
    st.caption("Write any SELECT query. The connection is read-only so INSERT/UPDATE/DELETE will be blocked.")

    sql = st.text_area(
        "SQL",
        height=150,
        placeholder="SELECT * FROM public.my_table LIMIT 100",
        key="sql_input",
    )

    col1, col2 = st.columns([1, 4])
    with col1:
        run_btn = st.button("Run Query", type="primary", key="run_sql")
    with col2:
        limit = st.number_input("Row limit", min_value=1, max_value=10000, value=500, step=100)

    if run_btn and sql.strip():
        if not is_read_only_query(sql):
            st.error("Only SELECT / read-only queries are allowed.")
        else:
            with st.spinner("Running..."):
                try:
                    df = run_query(**conn_params(), sql=sql, limit=limit)
                    if df.empty:
                        st.info("Query returned no rows.")
                    else:
                        st.success(f"Returned {len(df)} rows")
                        st.dataframe(df, use_container_width=True, hide_index=True)
                        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                        base = f"query_{ts}"
                        dc1, dc2 = st.columns(2)
                        with dc1:
                            st.download_button(
                                "Download CSV",
                                df_to_csv_bytes(df),
                                f"{base}.csv",
                                "text/csv",
                                key="sql_dl_csv",
                            )
                        with dc2:
                            st.download_button(
                                "Download Excel",
                                df_to_excel_bytes(df),
                                f"{base}.xlsx",
                                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                key="sql_dl_xlsx",
                            )
                except Exception as e:
                    st.error(f"Query error: {e}")

# --- AI Assistant ---
with tab_nl:
    st.header("AI Data Assistant")

    if not current_api_key():
        provider_name = "Anthropic" if st.session_state.provider == "anthropic" else "Gemini"
        link = "console.anthropic.com" if st.session_state.provider == "anthropic" else "aistudio.google.com/apikey"
        st.warning(f"Enter your {provider_name} API key in the sidebar to enable the AI assistant. Get one at [{link}](https://{link})")
        st.stop()

    # Show current provider/model/tier
    active_model = pick_model("agent")  # representative model for the banner
    st.caption(
        f"Ask anything about your data in plain English. "
        f"**Provider:** {st.session_state.provider.title()} · "
        f"**Model:** {st.session_state.model_label} · "
        f"**Tier:** {st.session_state.tier.title()}"
    )

    # Display chat history
    for msg_idx, msg in enumerate(st.session_state.chat_history):
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

            # Render agent-produced artifacts (Excel/Word files)
            for art_id in msg.get("artifact_ids", []):
                art = st.session_state.artifacts.get(art_id)
                if not art:
                    continue
                icon = "📊" if art["type"] == "excel" else "📄"
                with st.container(border=True):
                    st.markdown(f"**{icon} {art['title']}**")
                    if art.get("summary"):
                        st.caption(art["summary"])
                    st.caption(
                        f"{art['filename']} · {skills.format_size(art['size_bytes'])}"
                        + (f" · {art.get('sheet_count')} sheets" if art["type"] == "excel" else "")
                    )
                    mime = (
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        if art["type"] == "excel"
                        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.download_button(
                        "⬇ Download",
                        art["bytes"],
                        art["filename"],
                        mime,
                        key=f"dl_art_{art_id}",
                    )

            if msg.get("sql"):
                with st.expander("SQL query used"):
                    st.code(msg["sql"], language="sql")
            if msg.get("dataframe") is not None:
                df_msg = pd.DataFrame(msg["dataframe"])
                with st.expander(f"Raw data ({len(df_msg)} rows)"):
                    st.dataframe(df_msg, use_container_width=True, hide_index=True)

                # Find the corresponding user question (the preceding user message)
                user_question = ""
                for prev in st.session_state.chat_history[:msg_idx][::-1]:
                    if prev["role"] == "user":
                        user_question = prev["content"]
                        break

                # Export buttons
                st.markdown("**Export this result:**")
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                base = f"db_explorer_{ts}"
                c1, c2, c3, c4 = st.columns(4)
                with c1:
                    st.download_button(
                        "CSV (data)",
                        df_to_csv_bytes(df_msg),
                        f"{base}.csv",
                        "text/csv",
                        key=f"dl_csv_{msg_idx}",
                    )
                with c2:
                    st.download_button(
                        "Excel (data)",
                        df_to_excel_bytes(df_msg),
                        f"{base}.xlsx",
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key=f"dl_xlsx_{msg_idx}",
                    )
                with c3:
                    st.download_button(
                        "Word (full report)",
                        build_docx_report(user_question, msg["content"], msg.get("sql", ""), df_msg),
                        f"{base}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_docx_{msg_idx}",
                    )
                with c4:
                    st.download_button(
                        "Text (full report)",
                        build_text_report(user_question, msg["content"], msg.get("sql", ""), df_msg),
                        f"{base}.txt",
                        "text/plain",
                        key=f"dl_txt_{msg_idx}",
                    )

    # Clear chat button
    if st.session_state.chat_history:
        if st.button("Clear conversation", key="clear_chat"):
            st.session_state.chat_history = []
            st.session_state.artifacts = {}
            st.rerun()

    # Chat input
    if question := st.chat_input("Ask about your data... e.g., 'What are the top 10 tables by row count?'"):
        # Add user message to history
        st.session_state.chat_history.append({"role": "user", "content": question})
        with st.chat_message("user"):
            st.markdown(question)

        with st.chat_message("assistant"):
            schema_text = get_schema_summary(**conn_params())

            # Step 1: Generate SQL (or determine no SQL needed)
            with st.spinner("Thinking..."):
                sql, error = ask_claude_for_sql(
                    question, schema_text, current_api_key(),
                    chat_history=st.session_state.chat_history[:-1],
                    model=pick_model("sql"), tier=st.session_state.tier,
                )

            if error:
                response = f"Sorry, I ran into an error: {error}"
                st.markdown(response)
                st.session_state.chat_history.append({"role": "assistant", "content": response})

            elif sql and sql.strip().startswith("USE_AGENT_MODE"):
                # Multi-step agent with tools (Excel/Word artifacts, multi-query)
                status_placeholder = st.empty()

                def update_status(msg):
                    status_placeholder.info(f"🤖 {msg}")

                update_status("Starting agent...")
                response, new_artifacts, agent_err = run_agent_loop(
                    question, schema_text, current_api_key(),
                    conn_kwargs=conn_params(),
                    chat_history=st.session_state.chat_history[:-1],
                    progress_cb=update_status,
                    model=pick_model("agent"), tier=st.session_state.tier,
                )
                status_placeholder.empty()

                if agent_err:
                    response = f"Sorry, the agent ran into an error: {agent_err}"

                # Store artifacts in session state, keyed by id
                artifact_ids = []
                for art in new_artifacts:
                    st.session_state.artifacts[art["id"]] = art
                    artifact_ids.append(art["id"])

                st.markdown(response)
                # Show artifacts inline immediately (before rerun)
                for art in new_artifacts:
                    icon = "📊" if art["type"] == "excel" else "📄"
                    with st.container(border=True):
                        st.markdown(f"**{icon} {art['title']}**")
                        if art.get("summary"):
                            st.caption(art["summary"])
                        st.caption(
                            f"{art['filename']} · {skills.format_size(art['size_bytes'])}"
                        )

                st.session_state.chat_history.append({
                    "role": "assistant",
                    "content": response,
                    "artifact_ids": artifact_ids,
                })

            elif sql and sql.strip().startswith("NO_SQL_NEEDED"):
                # Conversational response, no query needed
                with st.spinner("Thinking..."):
                    response, error = ask_claude_no_sql(
                        question, schema_text, current_api_key(),
                        chat_history=st.session_state.chat_history[:-1],
                        model=pick_model("conversational"), tier=st.session_state.tier,
                    )
                if error:
                    response = f"Sorry, I ran into an error: {error}"
                st.markdown(response)
                st.session_state.chat_history.append({"role": "assistant", "content": response})

            elif sql:
                if not is_read_only_query(sql):
                    # AI returned text instead of SQL — the SQL-generation call caps at 1024
                    # tokens which can truncate long answers. Re-ask as a conversational query
                    # with the higher 4096-token limit to get a complete response.
                    logger.info(f"AI returned text instead of SQL, re-asking via chat path for full response")
                    with st.spinner("Generating full response..."):
                        response, chat_err = ask_claude_no_sql(
                            question, schema_text, current_api_key(),
                            chat_history=st.session_state.chat_history[:-1],
                            model=pick_model("conversational"), tier=st.session_state.tier,
                        )
                    if chat_err:
                        # Fall back to showing the (potentially truncated) original text
                        response = sql
                    st.markdown(response)
                    st.session_state.chat_history.append({"role": "assistant", "content": response})
                else:
                    # Step 2: Run the query (with auto-retry on failure)
                    max_retries = 2
                    current_sql = sql
                    df = None
                    last_error = None

                    for attempt in range(1, max_retries + 1):
                        with st.spinner("Fetching data..." if attempt == 1 else f"Retrying (attempt {attempt})..."):
                            try:
                                df = run_query(**conn_params(), sql=current_sql, limit=500)
                                last_error = None
                                break
                            except Exception as e:
                                last_error = str(e)
                                logger.warning(f"QUERY ATTEMPT {attempt} FAILED | {e}")
                                if attempt < max_retries:
                                    # Ask AI to fix the SQL
                                    with st.spinner("Fixing query..."):
                                        fix_prompt = f"Your previous SQL query had an error. Fix it.\n\nFailed SQL:\n{current_sql}\n\nError:\n{last_error}\n\nReturn ONLY the corrected SQL query, nothing else."
                                        fixed_sql, fix_err = ask_claude_for_sql(
                                            fix_prompt, schema_text, current_api_key(),
                                            model=pick_model("sql"), tier=st.session_state.tier,
                                        )
                                    if fix_err or not fixed_sql or not is_read_only_query(fixed_sql):
                                        logger.error(f"AI could not fix query: {fix_err}")
                                        break
                                    logger.info(f"AI RETRY SQL | {fixed_sql[:150]}")
                                    current_sql = fixed_sql

                    if last_error:
                        response = f"The query failed after {max_retries} attempts: {last_error}"
                        st.markdown(response)
                        st.session_state.chat_history.append({"role": "assistant", "content": response, "sql": sql})

                    with st.expander("SQL query used"):
                        st.code(current_sql, language="sql")

                    if df is not None:
                        if df.empty:
                            data_csv = "(no rows returned)"
                        else:
                            with st.expander(f"Raw data ({len(df)} rows)"):
                                st.dataframe(df, use_container_width=True, hide_index=True)
                            data_csv = df.head(100).to_csv(index=False)

                        # Step 3: Analyze results with AI
                        with st.spinner("Analyzing results..."):
                            analysis, error = ask_claude_to_analyze(
                                question, schema_text, data_csv, sql, current_api_key(),
                                chat_history=st.session_state.chat_history[:-1],
                                model=pick_model("analyze"), tier=st.session_state.tier,
                            )

                        if error:
                            analysis = f"I fetched the data but couldn't analyze it: {error}"

                        st.markdown(analysis)
                        st.session_state.chat_history.append({
                            "role": "assistant",
                            "content": analysis,
                            "summary": f"[Answered about: {question}]",
                            "sql": sql,
                            "dataframe": df.to_dict(orient="records") if not df.empty else None,
                        })

            else:
                # Fallback: no SQL and no error (shouldn't happen, but handle gracefully)
                logger.warning(f"AI returned unexpected response: sql={sql!r}, error={error!r}")
                response = "I couldn't understand that question. Could you try rephrasing it?"
                st.markdown(response)
                st.session_state.chat_history.append({"role": "assistant", "content": response})

        st.rerun()

# --- Quick Analytics ---
with tab_analytics:
    st.header("Quick Analytics")
    st.caption("Run a query in the SQL or English tab first, then come here to visualize.")

    analytics_sql = st.text_area(
        "SQL for chart",
        height=100,
        placeholder="SELECT category, SUM(amount) as total FROM orders GROUP BY category",
        key="analytics_sql",
    )

    if st.button("Run & Visualize", type="primary", key="run_analytics") and analytics_sql.strip():
        if not is_read_only_query(analytics_sql):
            st.error("Only SELECT queries are allowed.")
        else:
            with st.spinner("Running..."):
                try:
                    df = run_query(**conn_params(), sql=analytics_sql, limit=5000)
                    if df.empty:
                        st.info("No data returned.")
                    else:
                        st.dataframe(df, use_container_width=True, hide_index=True)

                        st.subheader("Chart Settings")
                        cols = list(df.columns)
                        c1, c2, c3 = st.columns(3)
                        with c1:
                            chart_type = st.selectbox("Chart type", ["Bar", "Line", "Scatter", "Pie", "Histogram"])
                        with c2:
                            x_col = st.selectbox("X axis", cols, index=0)
                        with c3:
                            y_col = st.selectbox("Y axis", cols, index=min(1, len(cols) - 1))

                        chart_map = {
                            "Bar": px.bar,
                            "Line": px.line,
                            "Scatter": px.scatter,
                            "Pie": px.pie,
                            "Histogram": px.histogram,
                        }
                        try:
                            if chart_type == "Pie":
                                fig = px.pie(df, names=x_col, values=y_col)
                            else:
                                fig = chart_map[chart_type](df, x=x_col, y=y_col)
                            st.plotly_chart(fig, use_container_width=True)
                        except Exception as e:
                            st.error(f"Chart error: {e}")
                except Exception as e:
                    st.error(f"Query error: {e}")

# --- Activity Log ---
with tab_logs:
    st.header("Activity Log")
    st.caption("Application logs for troubleshooting and review. Logs are saved to the `logs/` folder.")

    # List available log files
    log_files = sorted(LOG_DIR.glob("app_*.log"), reverse=True)

    if not log_files:
        st.info("No log files yet. Logs will appear here as you use the app.")
    else:
        selected_log = st.selectbox(
            "Log file",
            log_files,
            format_func=lambda f: f.name,
        )

        col1, col2, col3 = st.columns([1, 1, 2])
        with col1:
            show_level = st.selectbox("Filter level", ["ALL", "ERROR", "WARNING", "INFO"], index=0)
        with col2:
            tail_lines = st.number_input("Last N lines", min_value=10, max_value=5000, value=100, step=50)
        with col3:
            search_term = st.text_input("Search logs", placeholder="e.g. QUERY FAILED, AI ANALYZE")

        if st.button("Refresh", key="refresh_logs"):
            st.rerun()

        # Read and filter log content
        try:
            all_lines = selected_log.read_text(encoding="utf-8").splitlines()
            lines = all_lines[-tail_lines:]

            if show_level != "ALL":
                lines = [l for l in lines if f"| {show_level}" in l]

            if search_term.strip():
                lines = [l for l in lines if search_term.lower() in l.lower()]

            log_text = "\n".join(lines)

            if not lines:
                st.info("No matching log entries.")
            else:
                st.text_area(
                    f"{len(lines)} log entries",
                    value=log_text,
                    height=400,
                    key="log_display",
                )
                st.download_button(
                    "Download this log",
                    log_text,
                    file_name=selected_log.name,
                    mime="text/plain",
                )
        except Exception as e:
            st.error(f"Could not read log file: {e}")
