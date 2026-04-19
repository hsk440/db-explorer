"""
Skills = tools exposed to the AI agent for multi-step tasks.

Claude-inspired design:
- Each skill has a rich description so the model picks the right one
- Results are compact/truncated to keep token usage sane
- Artifact skills return bytes that get stored for download
"""

import io
import re
import uuid
import logging
from datetime import datetime
from typing import Any
import pandas as pd


logger = logging.getLogger("db_explorer")


# ---------------------------------------------------------------------------
# Tool schemas (sent to the Anthropic API)
# ---------------------------------------------------------------------------

TOOL_DEFINITIONS = [
    {
        "name": "query_database",
        "description": (
            "Execute a read-only PostgreSQL SELECT query against the user's database "
            "and return the results as JSON rows. Use this whenever you need to look "
            "up data from the database. You can call this multiple times to gather "
            "data across different queries. Results larger than 50 rows are truncated "
            "with a summary. If the query fails, the error message is returned so you "
            "can fix it and retry."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "The PostgreSQL SELECT query to execute. Must start with SELECT or WITH.",
                },
                "purpose": {
                    "type": "string",
                    "description": "Brief description of why you're running this query (shown to user).",
                },
            },
            "required": ["query", "purpose"],
        },
    },
    {
        "name": "create_excel_artifact",
        "description": (
            "Create a downloadable Excel (.xlsx) file with one or more named sheets. "
            "Use this when the user asks for data in an Excel file, spreadsheet, or "
            "workbook. You control the sheet organization — each key in `sheets` "
            "becomes a tab, and its value is a list of row objects (dicts with column "
            "name → value). After calling this, tell the user the file is ready below."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Filename ending in .xlsx, e.g. 'tables_by_category.xlsx'",
                },
                "title": {
                    "type": "string",
                    "description": "Short human-readable title for the artifact card.",
                },
                "summary": {
                    "type": "string",
                    "description": "One-line description of what's in the file.",
                },
                "sheets": {
                    "type": "object",
                    "description": (
                        "Map of sheet name -> list of row objects. "
                        "Example: {'Sheet1': [{'col_a': 1, 'col_b': 'x'}, {'col_a': 2, 'col_b': 'y'}]}"
                    ),
                },
            },
            "required": ["filename", "title", "summary", "sheets"],
        },
    },
    {
        "name": "create_word_artifact",
        "description": (
            "Create a downloadable Word (.docx) report with formatted sections, "
            "headings, paragraphs, and optional data tables. Use this when the user "
            "asks for a report, document, Word file, or written analysis. Each "
            "section has a heading, markdown-formatted body text, and optionally "
            "a data table."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Filename ending in .docx",
                },
                "title": {
                    "type": "string",
                    "description": "Document title (shown in the document and on the card).",
                },
                "summary": {
                    "type": "string",
                    "description": "One-line description for the artifact card.",
                },
                "sections": {
                    "type": "array",
                    "description": "Ordered list of document sections.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "heading": {"type": "string"},
                            "body": {
                                "type": "string",
                                "description": "Markdown text. Supports **bold**, bullet lists (- item), and headings (## subheading).",
                            },
                            "table_data": {
                                "type": "array",
                                "description": "Optional list of row objects to render as a formatted table after the body.",
                                "items": {"type": "object"},
                            },
                        },
                        "required": ["heading", "body"],
                    },
                },
            },
            "required": ["filename", "title", "summary", "sections"],
        },
    },
]


# ---------------------------------------------------------------------------
# Result truncation (token optimization)
# ---------------------------------------------------------------------------

MAX_ROWS_FULL = 50
MAX_ROWS_TRUNCATED_HEAD = 20
MAX_ROWS_TRUNCATED_TAIL = 5
MAX_RESULT_BYTES = 8_000


def _summarize_df(df: pd.DataFrame) -> dict:
    """Build a compact summary of a DataFrame for the AI."""
    summary = {"row_count": len(df), "columns": list(df.columns)}
    if not df.empty:
        dtypes = {col: str(df[col].dtype) for col in df.columns}
        summary["dtypes"] = dtypes
        # Null counts per column
        summary["null_counts"] = {col: int(df[col].isna().sum()) for col in df.columns}
    return summary


def truncate_query_result(df: pd.DataFrame) -> dict:
    """Convert a DataFrame to a compact dict structure for the AI tool response."""
    if df.empty:
        return {"status": "ok", "row_count": 0, "rows": [], "note": "No rows returned."}

    n = len(df)
    if n <= MAX_ROWS_FULL:
        rows = df.to_dict(orient="records")
        # Normalize NaN to None for JSON
        rows = [{k: (None if pd.isna(v) else v) for k, v in r.items()} for r in rows]
        payload = {"status": "ok", "row_count": n, "rows": rows}
    else:
        head = df.head(MAX_ROWS_TRUNCATED_HEAD).to_dict(orient="records")
        tail = df.tail(MAX_ROWS_TRUNCATED_TAIL).to_dict(orient="records")
        head = [{k: (None if pd.isna(v) else v) for k, v in r.items()} for r in head]
        tail = [{k: (None if pd.isna(v) else v) for k, v in r.items()} for r in tail]
        payload = {
            "status": "ok",
            "row_count": n,
            "rows_head": head,
            "rows_tail": tail,
            "summary": _summarize_df(df),
            "note": f"Showing first {MAX_ROWS_TRUNCATED_HEAD} and last {MAX_ROWS_TRUNCATED_TAIL} rows of {n} total.",
        }

    # Final byte-size clamp
    import json
    serialized = json.dumps(payload, default=str)
    if len(serialized) > MAX_RESULT_BYTES:
        payload = {
            "status": "ok",
            "row_count": n,
            "rows_head": df.head(5).to_dict(orient="records"),
            "summary": _summarize_df(df),
            "note": f"Result too large to embed fully. Showing first 5 of {n} rows + summary.",
        }

    return payload


# ---------------------------------------------------------------------------
# Artifact builders
# ---------------------------------------------------------------------------

def build_excel_artifact(filename: str, title: str, summary: str, sheets: dict) -> dict:
    """Build an Excel artifact from a dict of sheet_name -> list of row dicts."""
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        any_sheet = False
        for sheet_name, rows in sheets.items():
            safe_name = str(sheet_name)[:31] or "Sheet"
            if not rows:
                # Empty sheet with note
                pd.DataFrame({"info": ["No data in this section."]}).to_excel(
                    writer, sheet_name=safe_name, index=False
                )
            else:
                pd.DataFrame(rows).to_excel(writer, sheet_name=safe_name, index=False)
            any_sheet = True
        if not any_sheet:
            pd.DataFrame({"info": ["Empty workbook."]}).to_excel(
                writer, sheet_name="Sheet1", index=False
            )

    data = buf.getvalue()
    return {
        "id": str(uuid.uuid4()),
        "type": "excel",
        "filename": filename if filename.endswith(".xlsx") else f"{filename}.xlsx",
        "title": title,
        "summary": summary,
        "bytes": data,
        "size_bytes": len(data),
        "created_at": datetime.now().isoformat(),
        "sheet_count": len(sheets),
    }


def _add_markdown_runs(paragraph, text: str):
    """Add text to a docx paragraph, handling **bold** markdown."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def build_word_artifact(filename: str, title: str, summary: str, sections: list) -> dict:
    """Build a Word artifact with title, sections, and optional tables."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True

    for section in sections:
        heading = section.get("heading", "")
        body = section.get("body", "")
        table_data = section.get("table_data")

        if heading:
            doc.add_heading(heading, level=1)

        for line in body.split("\n"):
            line = line.rstrip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                _add_markdown_runs(p, line[2:])
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            else:
                p = doc.add_paragraph()
                _add_markdown_runs(p, line)

        if table_data:
            df = pd.DataFrame(table_data)
            if not df.empty:
                display_df = df.head(100)
                table = doc.add_table(rows=1, cols=len(display_df.columns))
                table.style = "Light Grid Accent 1"
                hdr = table.rows[0].cells
                for i, col in enumerate(display_df.columns):
                    hdr[i].text = str(col)
                    for para in hdr[i].paragraphs:
                        for run in para.runs:
                            run.bold = True
                for _, row in display_df.iterrows():
                    cells = table.add_row().cells
                    for i, val in enumerate(row):
                        cells[i].text = "" if pd.isna(val) else str(val)
                if len(df) > 100:
                    doc.add_paragraph(f"(showing first 100 of {len(df)} rows)")

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    return {
        "id": str(uuid.uuid4()),
        "type": "word",
        "filename": filename if filename.endswith(".docx") else f"{filename}.docx",
        "title": title,
        "summary": summary,
        "bytes": data,
        "size_bytes": len(data),
        "created_at": datetime.now().isoformat(),
        "section_count": len(sections),
    }


# ---------------------------------------------------------------------------
# Tool dispatcher
# ---------------------------------------------------------------------------

def dispatch_tool(tool_name: str, tool_input: dict, run_query_fn, is_safe_fn, conn_kwargs: dict):
    """
    Execute a tool call. Returns (result_payload, artifact_or_None).

    - tool_name: "query_database" | "create_excel_artifact" | "create_word_artifact"
    - tool_input: dict of the arguments the model provided
    - run_query_fn: function(**conn_kwargs, sql=, limit=) -> DataFrame
    - is_safe_fn: function(sql) -> bool (read-only check)
    - conn_kwargs: kwargs to pass to run_query_fn for DB connection
    """
    if tool_name == "query_database":
        query = tool_input.get("query", "").strip()
        purpose = tool_input.get("purpose", "")
        logger.info(f"AGENT query_database | purpose={purpose[:80]} | sql={query[:120]}")

        if not query:
            return {"status": "error", "error": "Empty query."}, None

        if not is_safe_fn(query):
            return {
                "status": "error",
                "error": "Only read-only SELECT / WITH / EXPLAIN / SHOW / VALUES queries are allowed.",
            }, None

        try:
            df = run_query_fn(**conn_kwargs, sql=query, limit=500)
            result = truncate_query_result(df)
            result["purpose"] = purpose
            return result, None
        except Exception as e:
            logger.warning(f"AGENT query_database FAILED | {e}")
            return {"status": "error", "error": str(e), "purpose": purpose}, None

    elif tool_name == "create_excel_artifact":
        try:
            artifact = build_excel_artifact(
                filename=tool_input.get("filename", "export.xlsx"),
                title=tool_input.get("title", "Excel Export"),
                summary=tool_input.get("summary", ""),
                sheets=tool_input.get("sheets", {}),
            )
            logger.info(f"AGENT created excel artifact | {artifact['filename']} | {artifact['size_bytes']} bytes")
            return {
                "status": "ok",
                "artifact_id": artifact["id"],
                "filename": artifact["filename"],
                "size_bytes": artifact["size_bytes"],
                "note": "Excel artifact created and will be shown to the user for download.",
            }, artifact
        except Exception as e:
            logger.error(f"AGENT create_excel FAILED | {e}")
            return {"status": "error", "error": str(e)}, None

    elif tool_name == "create_word_artifact":
        try:
            artifact = build_word_artifact(
                filename=tool_input.get("filename", "report.docx"),
                title=tool_input.get("title", "Report"),
                summary=tool_input.get("summary", ""),
                sections=tool_input.get("sections", []),
            )
            logger.info(f"AGENT created word artifact | {artifact['filename']} | {artifact['size_bytes']} bytes")
            return {
                "status": "ok",
                "artifact_id": artifact["id"],
                "filename": artifact["filename"],
                "size_bytes": artifact["size_bytes"],
                "note": "Word artifact created and will be shown to the user for download.",
            }, artifact
        except Exception as e:
            logger.error(f"AGENT create_word FAILED | {e}")
            return {"status": "error", "error": str(e)}, None

    else:
        return {"status": "error", "error": f"Unknown tool: {tool_name}"}, None


def format_size(n_bytes: int) -> str:
    """Human-friendly file size."""
    if n_bytes < 1024:
        return f"{n_bytes} B"
    if n_bytes < 1024 * 1024:
        return f"{n_bytes / 1024:.1f} KB"
    return f"{n_bytes / (1024 * 1024):.2f} MB"
